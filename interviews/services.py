import os
import logging
import requests
import time
from twilio.rest import Client
from twilio.twiml.voice_response import VoiceResponse
import PyPDF2
from docx import Document
import io
from django.conf import settings
from .models import Interview, InterviewResponse, InterviewResult
import json
import base64

# Set up loggers for different services
logger = logging.getLogger('interviews')
openai_logger = logging.getLogger('openai')
twilio_logger = logging.getLogger('twilio')

# Add comprehensive logging configuration
print(f"[DEBUG] Services: Initializing loggers")
print(f"[DEBUG] Services: OpenAI API Key set: {'YES' if os.getenv('OPENAI_API_KEY') else 'NO'}")
print(f"[DEBUG] Services: Twilio Account SID set: {'YES' if os.getenv('TWILIO_ACCOUNT_SID') else 'NO'}")
print(f"[DEBUG] Services: Twilio Auth Token set: {'YES' if os.getenv('TWILIO_AUTH_TOKEN') else 'NO'}")
print(f"[DEBUG] Services: Twilio Phone Number set: {'YES' if os.getenv('TWILIO_PHONE_NUMBER') else 'NO'}")
print(f"[DEBUG] Services: Webhook Base URL set: {'YES' if os.getenv('WEBHOOK_BASE_URL') else 'NO'}")

logger.info("Services: All service modules initialized")
openai_logger.info("OpenAI: OpenAI service logger initialized")
twilio_logger.info("Twilio: Twilio service logger initialized")

from openai import OpenAI
import openai

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
openai.api_key = os.getenv("OPENAI_API_KEY")



class OpenAIService:
    def __init__(self):
        self.model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        openai_logger.info(f"OpenAIService: Initialized with model {self.model}")

    def _make_request(self, prompt, temperature=0.7, max_tokens=500, max_retries=3):
        """Send prompt to OpenAI and return response text with retry logic"""
        for attempt in range(max_retries):
            try:
                openai_logger.info(
                    f"OpenAIService: Sending request to {self.model} (attempt {attempt+1}/{max_retries})"
                )
                start_time = time.time()
                response = client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=temperature,
                    max_tokens=max_tokens,
                )
                end_time = time.time()

                result = response.choices[0].message.content.strip()
                openai_logger.info(
                    f"OpenAIService: Request completed in {end_time - start_time:.2f}s"
                )
                return result
            except Exception as e:
                error_msg = str(e)
                openai_logger.warning(
                    f"OpenAIService: Error on attempt {attempt+1}/{max_retries}: {error_msg}"
                )
                # Retry only on rate-limit / server issues
                if any(x in error_msg for x in ["429", "503", "timeout"]):
                    if attempt < max_retries - 1:
                        wait_time = (attempt + 1) * 2
                        openai_logger.info(f"Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                        continue
                raise  # give up after last attempt or non-retryable error

    def clean_questions(self, raw_questions):
        cleaned = []
        for q in raw_questions:
            if not q or q.lower() == "json":
                continue
            q = q.strip().strip('"').strip("'")
            if q:
                cleaned.append(q)
        return cleaned

    def generate_questions_from_jd(self, job_title, job_description):
        """Generate 5–7 interview questions from job description"""
        openai_logger.info(f"Generating questions for job title: {job_title}")

        prompt = f"""
        Based on the following job description, generate 5–7 relevant interview questions.

        Job Title: {job_title}
        Job Description: {job_description}

        Generate questions that assess:
        1. Technical skills and experience
        2. Problem-solving abilities
        3. Communication skills
        4. Cultural fit
        5. Past achievements and challenges

        Return only the questions as a JSON array of strings, no additional text.
        """

        try:
            questions_text = self._make_request(prompt, temperature=0.7, max_tokens=500)
            # Remove code fences
            questions_text = questions_text.strip().strip("```").strip()

            try:
                questions = json.loads(questions_text)
                if not isinstance(questions, list):
                    raise ValueError("Not a list")
            except Exception:
                # fallback to line-based parsing
                lines = questions_text.splitlines()
                questions = [
                    line.strip().rstrip(",").strip('"').strip("'")
                    for line in lines
                    if line.strip() and line.strip() not in ["[", "]"]
                ]

            return self.clean_questions(questions)

        except Exception as e:
            openai_logger.error(
                f"Error generating questions for {job_title}: {str(e)}", exc_info=True
            )
            # fallback set
            fallback_questions = [
                "Can you tell me about your relevant experience for this role?",
                "What are your key strengths that would make you successful in this position?",
                "Describe a challenging project you worked on and how you handled it.",
                "What interests you most about this opportunity?",
                "Where do you see yourself professionally in the next few years?",
            ]
            return self.clean_questions(fallback_questions)

    def analyze_response(self, question, response_text, resume_context=""):
        prompt = f"""
        Analyze this interview response and provide a score (0–10) and feedback.

        Question: {question}
        Response: {response_text}
        Resume Context: {resume_context}

        Evaluate based on:
        - Relevance to the question
        - Clarity and communication
        - Specificity and examples
        - Professionalism

        Return as JSON: {{"score": float, "feedback": "string"}}
        """

        try:
            result_text = self._make_request(prompt, temperature=0.3, max_tokens=300)
            try:
                result = json.loads(result_text)
                return result.get("score", 5.0), result.get(
                    "feedback", "No feedback available."
                )
            except json.JSONDecodeError:
                return 5.0, "Analysis completed but format was unexpected."
        except Exception as e:
            openai_logger.error(f"Error analyzing response: {str(e)}", exc_info=True)
            return 5.0, "Unable to analyze response due to technical issues."

    def generate_final_recommendation(self, interview_responses, resume_context=""):
        """Generate final interview recommendation and overall score"""
        responses_summary = "\n".join(
            [
                f"Q{i+1}: {resp.question}\nA{i+1}: {resp.transcript}\nScore: {resp.score}"
                for i, resp in enumerate(interview_responses)
            ]
        )

        prompt = f"""
        Based on these interview responses, provide a comprehensive evaluation:

        Resume Context: {resume_context}

        Interview Responses:
        {responses_summary}

        Provide:
        1. Overall score (0–10)
        2. Recommendation (hire/consider/reject with reasoning)
        3. Key strengths (list)
        4. Areas for improvement (list)

        Return as JSON:
        {{
            "overall_score": float,
            "recommendation": "string",
            "strengths": ["string"],
            "areas_for_improvement": ["string"]
        }}
        """

        try:
            result_text = self._make_request(prompt, temperature=0.3, max_tokens=500)
            try:
                return json.loads(result_text)
            except json.JSONDecodeError:
                return {
                    "overall_score": 5.0,
                    "recommendation": "Consider - format unexpected",
                    "strengths": ["Analysis completed"],
                    "areas_for_improvement": ["Unable to parse response"],
                }
        except Exception as e:
            openai_logger.error(
                f"Error generating final recommendation: {str(e)}", exc_info=True
            )
            return {
                "overall_score": 5.0,
                "recommendation": "Consider - technical error",
                "strengths": ["Interview completed"],
                "areas_for_improvement": ["Technical analysis unavailable"],
            }
class TwilioService:
    """Service for Twilio voice call integration"""

    def __init__(self):
        self.client = Client(
            os.getenv("TWILIO_ACCOUNT_SID"),
            os.getenv("TWILIO_AUTH_TOKEN")
        )
        self.phone_number = os.getenv("TWILIO_PHONE_NUMBER")

        # Validate configuration
        if not os.getenv("TWILIO_ACCOUNT_SID"):
            twilio_logger.error("TwilioService: TWILIO_ACCOUNT_SID not set")
        if not os.getenv("TWILIO_AUTH_TOKEN"):
            twilio_logger.error("TwilioService: TWILIO_AUTH_TOKEN not set")
        if not self.phone_number:
            twilio_logger.error("TwilioService: TWILIO_PHONE_NUMBER not set")

    def initiate_call(self, interview_id, candidate_phone, questions):
        try:
            if not interview_id:
                raise ValueError("interview_id cannot be empty")
            if not candidate_phone:
                raise ValueError("candidate_phone cannot be empty")
            if not questions:
                raise ValueError("questions list cannot be empty")

            # Whitelist check
            whitelist_env = os.getenv("WHITELISTED_NUMBERS", '["*"]')
            try:
                whitelisted_numbers = json.loads(whitelist_env)
            except json.JSONDecodeError:
                whitelisted_numbers = ["*"]

            if "*" not in whitelisted_numbers and candidate_phone not in whitelisted_numbers:
                raise ValueError(f"Phone number {candidate_phone} is not whitelisted")

            # Webhook base URL
            webhook_base_url = os.getenv("WEBHOOK_BASE_URL", "http://localhost:8000")
            if not webhook_base_url.endswith("/"):
                webhook_base_url += "/"
            status_callback_url = f"{webhook_base_url}api/webhooks/call-status/"

            # TwiML for first question
            twiml = self._create_interview_twiml(interview_id, questions)

            # Create call
            call = self.client.calls.create(
                twiml=twiml,
                to=candidate_phone,
                from_=self.phone_number,
                record=True,
                status_callback=status_callback_url,
                status_callback_event=["completed"],
            )
            twilio_logger.info(f"TwilioService: Call initiated, SID={call.sid}")
            return call.sid

        except Exception as e:
            twilio_logger.error(f"TwilioService: Error initiating call: {str(e)}", exc_info=True)
            raise

    def _create_interview_twiml(self, interview_id, questions):
        """Generate TwiML for asking the first interview question"""
        try:
            response = VoiceResponse()

            if questions:
                response.say("Hello! Welcome to your automated interview. Let's begin.")
                response.pause(length=1)

                response.say(f"Question: {questions[0]}")
                response.pause(length=1)
                response.say("Please provide your answer now.")

                webhook_base_url = os.getenv("WEBHOOK_BASE_URL", "http://localhost:8000")
                if not webhook_base_url.endswith("/"):
                    webhook_base_url += "/"
                record_action_url = (
                    f"{webhook_base_url}api/webhooks/record-response/"
                    f"?interview_id={interview_id}&question_number=1"
                )

                response.record(
                    max_length=120,
                    play_beep=True,
                    action=record_action_url,
                    method="POST",
                    timeout=10,
                    transcribe=False,
                )
            else:
                response.say("No interview questions available. Goodbye!")
                response.hangup()

            return str(response)

        except Exception as e:
            twilio_logger.error(f"TwilioService: Error creating TwiML: {str(e)}", exc_info=True)
            fallback = VoiceResponse()
            fallback.say("We're experiencing technical difficulties. Please try again later.")
            fallback.hangup()
            return str(fallback)

class ResumeParserService:
    """Service for parsing resume files"""
    
    def __init__(self):
        print(f"[DEBUG] ResumeParserService: Initialized")
        logger.info("ResumeParserService: Initialized")
    
    def parse_resume(self, file):
        """Parse PDF or DOCX resume and extract text"""
        print(f"[DEBUG] ResumeParserService: Starting resume parsing")
        print(f"[DEBUG] ResumeParserService: File name: {file.name}")
        print(f"[DEBUG] ResumeParserService: File size: {file.size if hasattr(file, 'size') else 'Unknown'} bytes")
        print(f"[DEBUG] ResumeParserService: File content type: {getattr(file, 'content_type', 'Unknown')}")
        
        try:
            logger.info(f"ResumeParserService: Parsing resume file: {file.name}")
            
            if file.name.lower().endswith('.pdf'):
                print(f"[DEBUG] ResumeParserService: Detected PDF file, using PDF parser")
                result = self._parse_pdf(file)
                print(f"[DEBUG] ResumeParserService: PDF parsing completed, extracted {len(result)} characters")
                logger.info(f"ResumeParserService: Successfully parsed PDF resume, extracted {len(result)} characters")
                return result
            elif file.name.lower().endswith('.docx'):
                print(f"[DEBUG] ResumeParserService: Detected DOCX file, using DOCX parser")
                result = self._parse_docx(file)
                print(f"[DEBUG] ResumeParserService: DOCX parsing completed, extracted {len(result)} characters")
                logger.info(f"ResumeParserService: Successfully parsed DOCX resume, extracted {len(result)} characters")
                return result
            else:
                print(f"[ERROR] ResumeParserService: Unsupported file format: {file.name}")
                logger.error(f"ResumeParserService: Unsupported file format: {file.name}")
                raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
        except Exception as e:
            print(f"[ERROR] ResumeParserService: Error parsing resume {file.name}: {str(e)}")
            logger.error(f"ResumeParserService: Error parsing resume {file.name}: {str(e)}", exc_info=True)
            return "Unable to parse resume content."
     
    def _parse_pdf(self, file):
        try:
            print(f"[DEBUG] ResumeParserService: Starting PDF parsing for {file.name}")
            logger.info(f"ResumeParserService: Starting PDF parsing for {file.name}")
            
            pdf_reader = PyPDF2.PdfReader(file)
            print(f"[DEBUG] ResumeParserService: PDF reader created, number of pages: {len(pdf_reader.pages)}")
            
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                print(f"[DEBUG] ResumeParserService: Processing page {i+1}/{len(pdf_reader.pages)}")
                page_text = page.extract_text()
                print(f"[DEBUG] ResumeParserService: Page {i+1} extracted {len(page_text)} characters")
                text += page_text + "\n"
            
            result = text.strip()
            print(f"[DEBUG] ResumeParserService: PDF parsing completed")
            print(f"[DEBUG] ResumeParserService: Total extracted text length: {len(result)} characters")
            print(f"[DEBUG] ResumeParserService: Text preview: {result[:200]}...")
            
            logger.info(f"ResumeParserService: Successfully parsed PDF with {len(pdf_reader.pages)} pages")
            return result
        except Exception as e:
            print(f"[ERROR] ResumeParserService: Error parsing PDF {file.name}: {str(e)}")
            logger.error(f"ResumeParserService: Error parsing PDF {file.name}: {str(e)}", exc_info=True)
            return "Unable to extract text from PDF."
     
    def _parse_docx(self, file):
        try:
            print(f"[DEBUG] ResumeParserService: Starting DOCX parsing for {file.name}")
            logger.info(f"ResumeParserService: Starting DOCX parsing for {file.name}")
            
            doc = Document(file)
            print(f"[DEBUG] ResumeParserService: DOCX document loaded, number of paragraphs: {len(doc.paragraphs)}")
            
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                if i < 10:  # Log first 10 paragraphs for debugging
                    print(f"[DEBUG] ResumeParserService: Paragraph {i+1}: {len(paragraph.text)} characters")
                text += paragraph.text + "\n"
            
            result = text.strip()
            print(f"[DEBUG] ResumeParserService: DOCX parsing completed")
            print(f"[DEBUG] ResumeParserService: Total extracted text length: {len(result)} characters")
            print(f"[DEBUG] ResumeParserService: Text preview: {result[:200]}...")
            
            logger.info(f"ResumeParserService: Successfully parsed DOCX with {len(doc.paragraphs)} paragraphs")
            return result
        except Exception as e:
            print(f"[ERROR] ResumeParserService: Error parsing DOCX {file.name}: {str(e)}")
            logger.error(f"ResumeParserService: Error parsing DOCX {file.name}: {str(e)}", exc_info=True)
            return "Unable to extract text from DOCX."


class TranscriptionService:
    """Service for transcribing audio recordings"""

    def __init__(self):
        logger.info("TranscriptionService: Initializing")

        # OpenAI client
        self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        logger.debug("TranscriptionService: OpenAI client initialized")

        # Twilio client
        self.twilio_client = Client(
            os.getenv("TWILIO_ACCOUNT_SID"),
            os.getenv("TWILIO_AUTH_TOKEN")
        )
        logger.debug("TranscriptionService: Twilio client initialized")

    def transcribe_audio(self, audio_url: str) -> str:
        """Fetch recording from Twilio and transcribe with OpenAI Whisper"""
        try:
            logger.info(f"TranscriptionService: Starting transcription for {audio_url}")

            recording_sid = self._extract_recording_sid(audio_url)
            if not recording_sid:
                raise ValueError("Invalid recording SID")

            recording = self.twilio_client.recordings(recording_sid).fetch()

            # Wait until recording is completed
            waited = 0
            while getattr(recording, "status", "") != "completed" and waited < 30:
                time.sleep(5)
                waited += 5
                recording = self.twilio_client.recordings(recording_sid).fetch()

            if getattr(recording, "status", "") != "completed":
                raise Exception(f"Recording not completed (status={recording.status})")


            # Construct media URL
            media_url = f"https://api.twilio.com{recording.uri.replace('.json', '')}.mp3"

            # Download audio

            base_uri = recording.uri.replace('.json', '')
            if not base_uri.endswith('.mp3'):
                timeout=30
            )
            if response.status_code != 200:
                raise Exception(f"Failed to download audio (HTTP {response.status_code})")

            audio_data = response.content

            # Transcribe with Whisper
            transcript = self.openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=("audio.mp3", io.BytesIO(audio_data), "audio/mpeg"),
                response_format="text"
            )

            result = transcript.strip()
            logger.info(f"TranscriptionService: Transcription successful")
            return result

        except Exception as e:
            logger.error(f"TranscriptionService: Transcription failed: {e}", exc_info=True)
            return f"Transcription failed: {e}"

    def _extract_recording_sid(self, audio_url: str) -> str | None:
        """Extract Twilio recording SID from URL"""
        try:
            if "/Recordings/" in audio_url:
                return audio_url.split("/Recordings/")[-1].split("?")[0]
            if "/recordings/" in audio_url:
                return audio_url.split("/recordings/")[-1].split("?")[0]

            # fallback: last segment
            sid = audio_url.split("/")[-1].split("?")[0]
            return sid if sid.startswith("RE") else None

        except Exception:
            return None
